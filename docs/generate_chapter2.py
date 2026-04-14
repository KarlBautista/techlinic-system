from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

h = doc.add_heading('Table 1. The Conceptual Model of the Study', level=3)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Create 3-column table: Input | Process | Output
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
headers = ['Input', 'Process', 'Output']
for i, text in enumerate(headers):
    p = table.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

# Content row
row = table.add_row().cells

# --- INPUT CELL ---
input_cell = row[0]
input_cell.paragraphs[0].clear()

def add_section(cell, title, items, is_first=False):
    if not is_first:
        p = cell.add_paragraph()
        p.space_before = Pt(4)
    else:
        p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    for item in items:
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

add_section(input_cell, 'Knowledge Requirements:', [
    'Record Management',
    'Electronic Medical Record (EMR)',
    'Healthcare Data Visualization',
    'Web-Based Application Development',
    'Database Management',
    'Medical Terminologies',
    'User Experience (UX) Design in Healthcare',
    'Security in Digital Health Systems',
], is_first=True)

add_section(input_cell, 'Software Requirements:', [
    'Visual Studio Code (Code Editor)',
    'React 19 with Vite (Frontend Framework)',
    'Node.js with Express.js (Backend Server)',
    'Supabase (Database & Authentication)',
    'Tailwind CSS (Styling Framework)',
    'ApexCharts (Data Visualization Library)',
    'Git / GitHub (Version Control)',
])

add_section(input_cell, 'Hardware Requirements:', [
    'Processor (Intel Core i3 or equivalent)',
    'Random Access Memory (4 GB minimum)',
    'Stable Internet Connection',
])

# --- PROCESS CELL ---
process_cell = row[1]
process_cell.paragraphs[0].clear()

add_section(process_cell, 'System Design:', [
    'System Architecture Diagram',
    'Context Diagram',
    'Data Flow Diagram',
    'Entity-Relationship Diagram',
    'System Flowchart',
    'Activity Diagram',
    'Site Map',
], is_first=True)

add_section(process_cell, 'System Creation:', [
    'Patient Registration and Management',
    'Doctor Diagnosis and Medical Documentation',
    'Medicine Inventory Management',
    'Analytics and Graphical Reporting',
    'Notification and Alert System',
    'Personnel Management',
    'Authentication and Role-Based Access Control',
])

add_section(process_cell, 'Testing and Debugging:', [
    'Functional Testing',
    'Unit Testing',
    'Integration Testing',
    'User Acceptance Testing (UAT)',
    'Performance Testing',
    'Security Testing',
])

add_section(process_cell, 'Evaluation:', [
    'ISO 25010 Software Quality Model:',
    '  - Functional Suitability',
    '  - Performance Efficiency',
    '  - Compatibility',
    '  - Usability',
    '  - Security',
])

# --- OUTPUT CELL ---
output_cell = row[2]
output_cell.paragraphs[0].clear()
p = output_cell.paragraphs[0]
run = p.add_run('TechClinic: A Health Record and Analytics System for Technological University of the Philippines \u2013 Manila Campus')
run.font.name = 'Times New Roman'
run.font.size = Pt(10)

# Set column widths
for r in table.rows:
    r.cells[0].width = Inches(2.2)
    r.cells[1].width = Inches(2.6)
    r.cells[2].width = Inches(1.7)

# --- Explanation paragraphs ---
doc.add_paragraph()

explanations = [
    ('Input', 'The input consists of the resources and requirements needed to develop the TechClinic system. This includes domain knowledge such as record management, electronic medical records, healthcare data visualization, web-based application development, database management, medical terminologies, UX design in healthcare, and digital health security. Software tools include Visual Studio Code as the code editor, React 19 with Vite for the frontend, Node.js with Express.js for the backend server, Supabase for database and authentication services, Tailwind CSS for responsive styling, ApexCharts for data visualization, and Git/GitHub for version control. Hardware requirements include a processor (Intel Core i3 or equivalent), at least 4 GB of RAM, and a stable internet connection for accessing the cloud-based database.'),
    ('Process', 'The process involves designing the system through architectural diagrams including the system architecture, context diagram, data flow diagram, entity-relationship diagram, system flowchart, activity diagram, and site map. System creation encompasses the development of core modules: patient registration and management, doctor diagnosis and medical documentation, medicine inventory management, analytics and graphical reporting, notification and alert system, personnel management, and authentication with role-based access control. Testing and debugging include functional testing, unit testing, integration testing, user acceptance testing, performance testing, and security testing. The system is evaluated using the ISO 25010 software quality model across five criteria: functional suitability, performance efficiency, compatibility, usability, and security.'),
    ('Output', 'The final output is TechClinic: A Health Record and Analytics System for Technological University of the Philippines \u2013 Manila Campus \u2014 a fully functional web-based application that digitizes health record management, provides real-time analytics through graphical reports, manages pharmaceutical inventory, and supports role-based clinical workflows for the TUP Manila clinic.'),
]

for title, text in explanations:
    h2 = doc.add_heading(title, level=4)
    for run in h2.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

output_path = r'c:\Users\karlb\OneDrive\Desktop\Techclinic\docs\Chapter2-Conceptual-Model-UPDATED.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
