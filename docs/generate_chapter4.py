from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ---- Evaluation Results heading ----
h = doc.add_heading('Evaluation Results', level=2)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

doc.add_paragraph(
    'The system was evaluated by respondents using the ISO 25010 software quality model criteria. '
    'The evaluation instrument used a 4-point Likert scale, where 4 represents "Strongly Agree," '
    '3 represents "Agree," 2 represents "Disagree," and 1 represents "Strongly Disagree." '
    'The following table summarizes the evaluation results.'
)

# ---- Table 4.2 ----
criteria = [
    ('A. Functional Suitability', None),
    ('1. The system accurately records and stores student health information.', '___', '___'),
    ('2. The system correctly generates graphical representations based on the collected health data.', '___', '___'),
    ('3. The system properly manages the pharmaceutical inventory, including stock tracking and automatic deduction.', '___', '___'),
    ('4. The system appropriately restricts access based on user roles (Nurse and Doctor).', '___', '___'),
    ('5. The student self-registration portal accurately captures and stores patient information.', '___', '___'),
    ('B. Performance Efficiency', None),
    ('1. The system responds to user interactions (page loads, form submissions, searches) within an acceptable time.', '___', '___'),
    ('2. The analytics charts render and update promptly when changing timeframes.', '___', '___'),
    ('3. The auto-refresh feature operates without noticeable delays or performance degradation.', '___', '___'),
    ('C. Compatibility', None),
    ('1. The system functions correctly across different web browsers (Chrome, Firefox, Edge, Safari).', '___', '___'),
    ('2. The responsive design adapts appropriately to different screen sizes (desktop, tablet, mobile).', '___', '___'),
    ('3. The system can coexist and operate alongside other web applications without conflicts.', '___', '___'),
    ('D. Usability', None),
    ('1. The system interface is easy to navigate and understand for first-time users.', '___', '___'),
    ('2. The visual design (color scheme, layout, typography) is professional and appropriate for a medical system.', '___', '___'),
    ('3. The forms provide clear labels, appropriate input validation, and helpful feedback messages.', '___', '___'),
    ('4. The graphical representations (charts and graphs) are easy to interpret and provide meaningful insights.', '___', '___'),
    ('E. Security', None),
    ('1. The authentication system effectively prevents unauthorized access to the system.', '___', '___'),
    ('2. Role-based access control properly restricts features based on user roles.', '___', '___'),
    ('3. Sensitive health data is adequately protected from unauthorized viewing or modification.', '___', '___'),
]

# Create table with header
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0].cells
for i, text in enumerate(['Criteria', 'Mean', 'Interpretation']):
    p = hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

for row_data in criteria:
    row = table.add_row().cells
    if row_data[1] is None:
        # Section header row
        p = row[0].paragraphs[0]
        run = p.add_run(row_data[0])
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Merge mean and interpretation cells visually (leave blank)
        for c in [row[1], row[2]]:
            c.paragraphs[0].text = ''
    else:
        for i, val in enumerate(row_data):
            p = row[i].paragraphs[0]
            run = p.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            if i > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(4.5)
    row.cells[1].width = Inches(0.8)
    row.cells[2].width = Inches(1.2)

p = doc.add_paragraph()
run = p.add_run('Table 4.2. ISO 25010 Evaluation Criteria and Results')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Note: The mean scores and interpretations in Table 4.2 are to be filled in after conducting '
    'the actual evaluation with the respondents. The interpretation follows the scale: '
    '3.50 – 4.00 = Strongly Agree, 2.50 – 3.49 = Agree, 1.50 – 2.49 = Disagree, '
    '1.00 – 1.49 = Strongly Disagree.'
)

# ---- Summary Table 4.3 ----
h2 = doc.add_heading('Summary of Evaluation Results', level=2)
for run in h2.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

summary_rows = [
    ('A. Functional Suitability', '___', '___'),
    ('B. Performance Efficiency', '___', '___'),
    ('C. Compatibility', '___', '___'),
    ('D. Usability', '___', '___'),
    ('E. Security', '___', '___'),
    ('Overall Mean', '___', '___'),
]

summary_table = doc.add_table(rows=1, cols=3)
summary_table.style = 'Table Grid'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr2 = summary_table.rows[0].cells
for i, text in enumerate(['Criteria', 'Mean', 'Interpretation']):
    p = hdr2[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

for row_data in summary_rows:
    row = summary_table.add_row().cells
    for i, val in enumerate(row_data):
        p = row[i].paragraphs[0]
        run = p.add_run(val)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        if i > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row_data[0] == 'Overall Mean':
            run.bold = True

for row in summary_table.rows:
    row.cells[0].width = Inches(3.0)
    row.cells[1].width = Inches(1.0)
    row.cells[2].width = Inches(1.5)

p = doc.add_paragraph()
run = p.add_run('Table 4.3. Summary of Evaluation Results')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---- Discussion ----
h3 = doc.add_heading('Discussion', level=2)
for run in h3.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

discussion_paragraphs = [
    'The results of the system testing demonstrate that all 30 test cases passed successfully, indicating that TechClinic meets its functional requirements as outlined in Chapter 1. The student self-registration portal, health record management, doctor diagnosis workflow, medicine inventory, analytics, notification system, personnel management, and settings functionalities all operate as intended.',
    'The system addresses the primary problem identified in the study \u2014 the reliance on manual tally sheets and paper-based record keeping at the TUP Manila clinic. By digitizing the health record management process, TechClinic eliminates the risk of data loss, reduces time spent on manual compilation, and provides immediate access to historical patient data.',
    'The analytics module directly fulfills the study\u2019s objective of generating graphical representations across different time periods. The four chart types \u2014 patient visit trends (area chart), department distribution (donut chart), top diagnoses (Pareto chart), and medicine stock levels (horizontal bar chart) \u2014 each support weekly, monthly, quarterly, and yearly views, enabling the medical staff to identify patterns and make informed decisions about health interventions.',
    'The automated notification system provides a proactive approach to disease monitoring that was not possible with the previous manual system. By automatically detecting when disease cases exceed threshold percentages, the system supports the study\u2019s objective of notifying medical staff of conditions that may necessitate adjustments to class modules or schedules.',
    'The role-based access control system ensures that the appropriate personnel perform their designated tasks \u2014 nurses handle patient intake while doctors manage diagnoses \u2014 maintaining a clear workflow that mirrors the actual clinic operations at TUP Manila. The three-layer security model (database RLS, backend middleware, frontend guards) provides defense in depth against unauthorized access.',
    'The evaluation results based on the ISO 25010 criteria will provide quantitative validation of the system\u2019s quality across five dimensions: functional suitability, performance efficiency, compatibility, usability, and security. These results, once gathered from the respondents, will offer a comprehensive assessment of how well TechClinic meets its design objectives compared to the traditional record-keeping methods it replaces.',
]

for text in discussion_paragraphs:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.5)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

output_path = r'c:\Users\karlb\OneDrive\Desktop\Techclinic\docs\Chapter4-Evaluation-Results-UPDATED.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
